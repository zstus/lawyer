"""DOCX 파일 파서 - 대출약정서에서 조/항 구조를 추출

설계 철학:
- 조(Article)와 항(Clause) 패턴을 유일한 분류 기준으로 사용
- 첫 번째 조 이전 = 문서 헤더/기본정보 (자동 추출)
- 첫 번째 조 이후 = 본문 (조/항 구조 파싱)
- 부록/별첨 = 파싱 종료 지점
"""
import re
from typing import List, Tuple, Optional, Any
from docx import Document
from .schemas import ParsedDocument, ParsedArticle, ParsedClause


# 조(Article) 패턴 매칭
# 예: "제 1 조	정	의", "제1조 대출약정", "제4조의2 시장붕괴", "제1조 정의 및 해석 2"
# 공백 선택적 ([\s\t]*), 제목 길이 제한 완화 (80자), 페이지 번호는 후처리로 제거
# 제목이 "제N항"으로 시작하면 안됨 (본문 내 참조와 구분)
ARTICLE_PATTERN = re.compile(
    r'^제[\s\t]*(\d+)[\s\t]*조(?:의[\s\t]*(\d+))?[\s\t]+(?!제[\s\t]*\d+[\s\t]*항)(.{1,80})$'
)

# 항(Clause) 패턴 매칭
# 예: "제 1 항	정	의", "제1항 차입의 종류", "제1항차입의 종류"
# 공백 선택적, 제목 길이 제한 완화 (100자)
CLAUSE_PATTERN = re.compile(
    r'^제[\s\t]*(\d+)[\s\t]*항[\s\t]*(.{1,100})$'
)

# 부록/별첨 패턴 (파싱 종료 지점)
# "부록 I", "부록 Ⅰ", "별첨 가-1" 형태로 시작하는 제목
APPENDIX_PATTERN = re.compile(
    r'^(부록|별첨|별지|첨부)[\s\t]*[IⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ\d가-힣\-]+[\s\t]*'
)

# 부록 참조 패턴 (본문 내 "부록 Ⅰ에" 같은 참조)
APPENDIX_REFERENCE_PATTERN = re.compile(
    r'^(부록|별첨|별지|첨부)[\s\t]*[IⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ\d가-힣\-]+[에의을를]'
)


def extract_text_from_docx(file_path: str) -> List[str]:
    """DOCX 파일에서 모든 문단의 텍스트를 추출"""
    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return paragraphs


def normalize_whitespace(text: str) -> str:
    """탭과 연속 공백을 단일 공백으로 정규화"""
    return re.sub(r'[\s\t]+', ' ', text).strip()


def clean_title(text: str) -> str:
    """제목에서 페이지 번호 및 불필요한 문자 제거

    예시:
    - "정의 및 해석 2" → "정의 및 해석"
    - "정의2" → "정의"
    - "대출 약정   15" → "대출 약정"
    """
    # 공백 정규화
    text = re.sub(r'[\s\t]+', ' ', text).strip()

    # 끝의 페이지 번호 제거 (공백 유무 관계없이)
    # 1-3자리 숫자 (페이지 번호는 보통 1-999)
    text = re.sub(r'\s+\d{1,3}$', '', text)
    # 한글 바로 뒤 숫자 (공백 없는 경우: "정의2")
    text = re.sub(r'([가-힣])\d{1,2}$', r'\1', text)

    return text.strip()


def find_first_article_index(paragraphs: List[str]) -> int:
    """본문의 첫 번째 조(Article) 헤더 인덱스를 찾음

    목차와 본문을 구분하는 방법:
    - 목차: 조 → 항 → 항 → 조 (연속된 헤더만 있음)
    - 본문: 조 → 항 → 내용 (항 다음에 실제 내용이 옴)

    Returns:
        본문의 첫 번째 조 인덱스, 없으면 0
    """
    for i, para in enumerate(paragraphs):
        if is_article_header(para):
            # 다음 4개 문단의 패턴 분석
            pattern = []
            for j in range(i, min(i + 5, len(paragraphs))):
                p = paragraphs[j]
                if is_article_header(p):
                    pattern.append('A')  # Article
                elif is_clause_header(p):
                    pattern.append('C')  # Clause
                else:
                    pattern.append('T')  # Text (내용)

            # 본문 패턴: 조 → 항 → 내용 (A-C-T)
            # 처음 3개 중에 T(내용)가 있으면 본문으로 판단
            if len(pattern) >= 3:
                # 조 다음에 항이 오고, 그 다음 2개 안에 내용이 있으면 본문
                if pattern[0] == 'A' and pattern[1] == 'C' and 'T' in pattern[2:4]:
                    return i
                # 조 다음에 바로 내용이 오면 본문 (항 없는 조)
                elif pattern[0] == 'A' and pattern[1] == 'T':
                    return i

    # 패턴을 찾지 못하면 첫 번째 조 반환
    for i, para in enumerate(paragraphs):
        if is_article_header(para):
            return i
    return 0


def extract_header_info(paragraphs: List[str], first_article_idx: int) -> dict:
    """문서 헤더(첫 번째 조 이전)에서 기본 정보 추출

    Returns:
        {"name": 문서명, "header_text": 헤더 전체 텍스트}
    """
    header_paragraphs = paragraphs[:first_article_idx]

    doc_name = None
    for para in header_paragraphs:
        # 공백 제거하여 키워드 검색
        normalized = para.replace(" ", "").replace("\t", "")
        if any(kw in normalized for kw in ["대출약정서", "약정서", "대출약정"]):
            # 공백 정규화하여 문서명 저장
            doc_name = normalize_whitespace(para)
            break

    return {
        "name": doc_name,
        "header_text": "\n".join(header_paragraphs) if header_paragraphs else ""
    }


def is_article_header(text: str) -> Optional[Tuple[int, Optional[int], str]]:
    """조 헤더인지 확인하고 (조번호, 조의번호, 타이틀) 반환"""
    match = ARTICLE_PATTERN.match(text)
    if match:
        article_num = int(match.group(1))
        sub_num = int(match.group(2)) if match.group(2) else None
        # 타이틀에서 탭과 연속 공백 정규화 및 페이지 번호 제거
        title = clean_title(match.group(3))
        # 빈 타이틀 체크 (제목 없는 조는 허용하지 않음)
        if not title:
            return None
        return (article_num, sub_num, title)
    return None


def is_clause_header(text: str) -> Optional[Tuple[int, str]]:
    """항 헤더인지 확인하고 (항번호, 타이틀) 반환"""
    match = CLAUSE_PATTERN.match(text)
    if match:
        clause_num = int(match.group(1))
        # 타이틀에서 탭과 연속 공백 정규화 및 페이지 번호 제거
        title = clean_title(match.group(2))
        # 빈 타이틀 체크 (제목 없는 항은 허용하지 않음)
        if not title:
            return None
        return (clause_num, title)
    return None


def is_appendix_start(text: str) -> bool:
    """부록/별첨 시작인지 확인 (본문 내 참조는 제외)"""
    # 본문 내 참조인지 먼저 확인 (예: "부록 Ⅰ에 기재된...")
    if APPENDIX_REFERENCE_PATTERN.match(text):
        return False
    return bool(APPENDIX_PATTERN.match(text))


def format_article_display(article_num: int, sub_num: Optional[int]) -> str:
    """조 번호를 표시용 문자열로 변환

    Returns:
        단순 번호 (예: "1", "4의2")
        템플릿에서 "제{display}조"로 표시됨
    """
    if sub_num:
        return f"{article_num}의{sub_num}"
    return f"{article_num}"


def parse_docx(file_path: str, file_name: str) -> ParsedDocument:
    """
    DOCX 파일을 파싱하여 조/항 구조를 추출

    일반화된 파싱 로직:
    1. 첫 번째 조 위치를 찾아 본문 시작점으로 사용
    2. 첫 번째 조 이전은 헤더로 처리하여 문서명 추출
    3. 첫 번째 조부터 부록/별첨 전까지 본문 파싱

    Args:
        file_path: DOCX 파일 경로
        file_name: 원본 파일명

    Returns:
        ParsedDocument: 파싱된 문서 구조
    """
    paragraphs = extract_text_from_docx(file_path)

    # 공통 파싱 로직 사용
    return _parse_paragraphs(paragraphs, file_name)


def _format_content(lines: List[str]) -> Any:
    """
    내용을 JSON 형태로 저장
    줄바꿈과 구조를 유지하면서 저장
    """
    if not lines:
        return None

    # 각 줄을 개별 요소로 저장하여 구조 유지
    return {
        "type": "clause_content",
        "lines": lines,
        "text": "\n".join(lines)  # 전체 텍스트도 저장
    }


def parse_docx_bytes(file_content: bytes, file_name: str) -> ParsedDocument:
    """
    바이트 데이터에서 DOCX 파싱 (업로드된 파일 처리용)

    Args:
        file_content: 파일 바이트 데이터
        file_name: 원본 파일명

    Returns:
        ParsedDocument: 파싱된 문서 구조
    """
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_content))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return _parse_paragraphs(paragraphs, file_name)


def _parse_paragraphs(paragraphs: List[str], file_name: str) -> ParsedDocument:
    """문단 리스트에서 조/항 구조 파싱

    일반화된 파싱 로직:
    1. 첫 번째 조 위치를 찾아 본문 시작점으로 사용
    2. 첫 번째 조 이전은 헤더로 처리하여 문서명 추출
    3. 첫 번째 조부터 부록/별첨 전까지 본문 파싱
    4. 항이 없는 조는 "본문" 항으로 처리
    """
    articles: List[ParsedArticle] = []
    current_article: Optional[ParsedArticle] = None
    current_clause_data: Optional[dict] = None
    clause_content_lines: List[str] = []
    article_content_lines: List[str] = []  # 항이 없는 조의 본문

    # 1. 첫 번째 조 위치 찾기 (본문 시작점)
    first_article_idx = find_first_article_index(paragraphs)

    # 2. 헤더 정보 추출 (첫 번째 조 이전)
    header_info = extract_header_info(paragraphs, first_article_idx)
    doc_name = header_info["name"] or file_name

    article_order = 0
    clause_order = 0
    in_main_content = False

    # 3. 본문 파싱 (첫 번째 조부터 시작)
    for para in paragraphs[first_article_idx:]:
        # 부록/별첨 시작시 종료
        if is_appendix_start(para):
            break

        # 조 헤더 확인
        article_info = is_article_header(para)
        if article_info:
            article_num, sub_num, title = article_info

            # 이전 항 내용 저장
            if current_clause_data is not None and clause_content_lines:
                _save_clause_content(articles, current_clause_data, clause_content_lines)
                clause_content_lines = []

            # 이전 조에 항이 없고 본문만 있는 경우 처리
            if current_article is not None and len(current_article.clauses) == 0 and article_content_lines:
                # "본문" 항 생성
                current_article.clauses.append(
                    ParsedClause(
                        clause_number=0,
                        clause_number_display="본문",
                        title="본문",
                        content=_format_content(article_content_lines),
                        order_index=1
                    )
                )
                article_content_lines = []

            article_order += 1
            current_article = ParsedArticle(
                article_number=article_num,
                article_number_display=format_article_display(article_num, sub_num),
                title=title,
                order_index=article_order,
                clauses=[]
            )
            articles.append(current_article)
            current_clause_data = None
            clause_order = 0
            article_content_lines = []  # 새 조 시작시 리셋
            in_main_content = True
            continue

        # 항 헤더 확인
        clause_info = is_clause_header(para)
        if clause_info and current_article:
            clause_num, title = clause_info

            # 이전 항 내용 저장
            if current_clause_data is not None and clause_content_lines:
                _save_clause_content(articles, current_clause_data, clause_content_lines)
                clause_content_lines = []

            # 항 헤더가 나왔으므로 조 본문 내용은 무시 (있었다면)
            article_content_lines = []

            clause_order += 1
            current_clause_data = {
                "article_idx": len(articles) - 1,
                "clause_idx": clause_order - 1
            }

            current_article.clauses.append(
                ParsedClause(
                    clause_number=clause_num,
                    clause_number_display=str(clause_num),
                    title=title,
                    content=None,
                    order_index=clause_order
                )
            )
            continue

        # 내용 수집
        if in_main_content:
            if current_clause_data is not None:
                # 현재 항에 내용 추가
                clause_content_lines.append(para)
            elif current_article is not None:
                # 항이 없는 조의 본문 내용 추가
                article_content_lines.append(para)

    # 마지막 항 내용 저장
    if current_clause_data is not None and clause_content_lines:
        _save_clause_content(articles, current_clause_data, clause_content_lines)

    # 마지막 조에 항이 없고 본문만 있는 경우 처리
    if current_article is not None and len(current_article.clauses) == 0 and article_content_lines:
        current_article.clauses.append(
            ParsedClause(
                clause_number=0,
                clause_number_display="본문",
                title="본문",
                content=_format_content(article_content_lines),
                order_index=1
            )
        )

    return ParsedDocument(
        name=doc_name,
        file_name=file_name,
        description=f"총 {len(articles)}개 조항",
        articles=articles
    )


def _save_clause_content(articles: List[ParsedArticle], clause_data: dict, lines: List[str]):
    """항 내용을 저장"""
    if not articles or not lines:
        return

    article_idx = clause_data.get("article_idx")
    clause_idx = clause_data.get("clause_idx")

    if article_idx is not None and clause_idx is not None:
        if article_idx < len(articles):
            article = articles[article_idx]
            if clause_idx < len(article.clauses):
                old_clause = article.clauses[clause_idx]
                article.clauses[clause_idx] = ParsedClause(
                    clause_number=old_clause.clause_number,
                    clause_number_display=old_clause.clause_number_display,
                    title=old_clause.title,
                    content=_format_content(lines),
                    order_index=old_clause.order_index
                )
